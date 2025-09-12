import io
import os
import tempfile

from docx import Document

from auris_tools.officeWordHandler import OfficeWordHandler
from auris_tools.storageHandler import StorageHandler

DOCX_SAMPLE = os.path.join(
    os.path.dirname(__file__), 'data', 'plain_text_sample.docx'
)
DOCX_SAMPLE_S3_BUCKET = 'auris-code-tests'
DOCX_SAMPLE_S3_KEY = 'plain_text_sample.docx'


def test_init_with_parameters():
    """Test initialization with parameters."""

    word_handler = OfficeWordHandler()
    assert isinstance(word_handler, OfficeWordHandler)
    assert word_handler is not None


def test_read_from_s3():
    """Test text extraction from a DOCX file."""

    word_handler = OfficeWordHandler()
    text = word_handler.read_from_s3(
        DOCX_SAMPLE_S3_BUCKET, DOCX_SAMPLE_S3_KEY, as_bytes_io=True
    )
    assert text is not None
    assert isinstance(text, io.BytesIO)


def test_read_from_s3_error():
    """Test error handling when reading a non-existent DOCX file from S3."""
    word_handler = OfficeWordHandler()
    try:
        word_handler.read_from_s3(
            'non_existent_bucket', 'non_existent_key.docx'
        )
    except Exception as e:
        assert isinstance(e, Exception)
        assert 'Error reading file from S3' in str(e)


def test_get_text_from_bytes():
    """Test text extraction from a DOCX file."""
    word_handler = OfficeWordHandler()
    text = word_handler.read_from_s3(
        DOCX_SAMPLE_S3_BUCKET, DOCX_SAMPLE_S3_KEY, as_bytes_io=True
    )
    # Extract text from the DOCX bytes
    extracted_text = word_handler.get_text_from_bytes(text.read())
    assert extracted_text is not None
    assert isinstance(extracted_text, str)
    assert len(extracted_text) > 0


def test_upload_to_s3():
    """Test uploading a DOCX file to S3."""
    word_handler = OfficeWordHandler()
    storage_handler = StorageHandler()

    # Create a temporary DOCX file

    # Create a simple document
    doc = Document()
    doc.add_paragraph(
        'This is a test document created for testing upload functionality.'
    )

    # Save to temporary file
    with tempfile.NamedTemporaryFile(
        suffix='.docx', delete=False
    ) as temp_file:
        temp_path = temp_file.name
        doc.save(temp_path)

    # Read the temporary DOCX file as a Document object
    with open(temp_path, 'rb') as f:
        docx_bytes = f.read()
        doc_obj = Document(io.BytesIO(docx_bytes))

    # Clean up the temporary file
    os.unlink(temp_path)

    # Upload the DOCX file to S3
    upload_key = 'test_upload_plain_text_sample.docx'
    word_handler.upload_docx(doc_obj, DOCX_SAMPLE_S3_BUCKET, upload_key)

    # Verify the file was uploaded
    s3_object = storage_handler.get_file_object(
        DOCX_SAMPLE_S3_BUCKET, upload_key, as_bytes=True
    )
    assert s3_object is not None
    assert s3_object == docx_bytes

    # Clean up the uploaded file from S3
    storage_handler.delete_file(DOCX_SAMPLE_S3_BUCKET, upload_key)


def test_upload_docx_error():
    """Test error handling when uploading a DOCX file to a non-existent S3 bucket."""
    word_handler = OfficeWordHandler()

    # Create a simple document
    doc = Document()
    doc.add_paragraph(
        'This is a test document created for testing upload functionality.'
    )

    # Save to temporary file
    with tempfile.NamedTemporaryFile(
        suffix='.docx', delete=False
    ) as temp_file:
        temp_path = temp_file.name
        doc.save(temp_path)

    # Read the temporary DOCX file as a Document object
    with open(temp_path, 'rb') as f:
        docx_bytes = f.read()
        doc_obj = Document(io.BytesIO(docx_bytes))

    # Clean up the temporary file
    os.unlink(temp_path)

    try:
        word_handler.upload_docx(doc_obj, 'non_existent_bucket', 'test.docx')
    except Exception as e:
        assert isinstance(e, Exception)
        assert 'Error uploading file to S3' in str(e)


def test_clean_text():
    """Test text cleaning functionality."""
    word_handler = OfficeWordHandler()
    raw_text = 'This is a test.\n\nThis is the second paragraph.\n\n\nThis is the third paragraph.'
    cleaned_text = word_handler.clean_text(raw_text)
    assert cleaned_text is not None


def test_clean_text_empty():
    """Test text cleaning with empty input."""
    word_handler = OfficeWordHandler()
    raw_text = ''
    cleaned_text = word_handler.clean_text(raw_text)
    assert cleaned_text == ''


def test_collect_all_paragraphs():
    """Test paragraph collection from a DOCX file."""
    word_handler = OfficeWordHandler()
    with open(DOCX_SAMPLE, 'rb') as f:
        docx_bytes = f.read()
        docx = Document(io.BytesIO(docx_bytes))

    paragraphs = word_handler.collect_all_paragraphs(docx)
    assert paragraphs is not None
    assert isinstance(paragraphs, list)
    assert len(paragraphs) > 0


def test_replace_placeholder_by_text_success():
    """Test replacing placeholder text in a DOCX file."""
    word_handler = OfficeWordHandler()
    with open(DOCX_SAMPLE, 'rb') as f:
        docx_bytes = f.read()
        docx = Document(io.BytesIO(docx_bytes))

    paragraphs = word_handler.collect_all_paragraphs(docx)
    placeholder = 'Lorem'
    replacement = 'BIRULEI'

    replacements_made = word_handler.replace_placeholder_by_text(
        paragraphs, docx, placeholder, replacement, max_count=2
    )

    assert replacements_made > 0

    # Verify that the replacements were made
    all_text = '\n'.join([p.text for p in paragraphs])
    assert replacement in all_text
    assert placeholder not in all_text
